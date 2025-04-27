from flask import Flask, request, jsonify, render_template, send_from_directory
import os
import json
from C1 import C4ContextDiagram
from C2 import C4ContainerDiagram
from C3 import C4ComponentDiagram
from C4 import C4CodeDiagram
from docx import Document

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
DIAGRAM_FOLDER = 'diagrams_output'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DIAGRAM_FOLDER, exist_ok=True)

def parse_docx_to_c4_json(filepath):
    doc = Document(filepath)
    system_name = "Unnamed System"
    users, external_systems, containers, components, relationships = [], [], [], [], []

    print("--- Parsing DOCX ---")

    # Extract system name (simplified)
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and para.text.strip():
            system_name = para.text.strip()
            print("  Found system name:", system_name)
            break
        elif para.text.strip():
            system_name = para.text.strip()
            print("  Found system name (from first para):", system_name)
            break

    for table_idx, table in enumerate(doc.tables):
        print(f"--- Table {table_idx + 1} ---")

        # Attempt to identify table type by content (customize this!)
        table_type = None
        if table_idx == 0:  # Assuming the first table is users/systems
            table_type = "users_systems"
        elif table_idx == 1:  # Assuming the second table is containers
            table_type = "containers"
        elif table_idx == 2:  # Assuming the third table is components
            table_type = "components"
        elif table_idx == 3:  # Assuming the fourth table is relationships
            table_type = "relationships"

        if not table_type:
            print("  Skipping table: Unable to determine type")
            continue  # Skip if table type is unknown

        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        print("  Table headers:", headers)

        if table_type == "users_systems" and "name" in headers and "type" in headers and "description" in headers:
            name_idx = headers.index("name")
            type_idx = headers.index("type")
            desc_idx = headers.index("description")
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3:
                    if cells[type_idx].lower() == "person":
                        users.append({"name": cells[name_idx], "description": cells[desc_idx]})
                        print("   Added user:", users[-1])
                    elif cells[type_idx].lower() == "system":
                        external_systems.append({"name": cells[name_idx], "description": cells[desc_idx]})
                        print("   Added external system:", external_systems[-1])

        elif table_type == "containers" and "container name" in headers and "technology" in headers:
            name_idx = headers.index("container name")
            tech_idx = headers.index("technology")
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    containers.append({"name": cells[name_idx], "technology": cells[tech_idx]})
                    print("   Added container:", containers[-1])

        elif table_type == "components" and "component name" in headers and "technology" in headers:
            name_idx = headers.index("component name")
            tech_idx = headers.index("technology")
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    components.append({"name": cells[name_idx], "technology": cells[tech_idx]})
                    print("   Added component:", components[-1])

        elif table_type == "relationships" and "from" in headers and "to" in headers and "label" in headers:
            from_idx = headers.index("from")
            to_idx = headers.index("to")
            label_idx = headers.index("label")
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3:
                    relationships.append({"source": cells[from_idx], "target": cells[to_idx], "label": cells[label_idx]})
                    print("   Added relationship:", relationships[-1])

    print("--- Parsing Complete ---")
    print("Extracted data:")
    print("System Name:", system_name)
    print("Users:", users)
    print("External Systems:", external_systems)
    print("Containers:", containers)
    print("Components:", components)
    print("Relationships:", relationships)

    return {
        "system_name": system_name,
        "users": users,
        "external_systems": external_systems,
        "containers": containers,
        "components": components,
        "relationships": relationships
    }

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    try:
        print("--- /upload route called ---")  # Entry point
        if 'file' not in request.files:
            print("Error: No file uploaded.")
            return jsonify({'error': 'No file uploaded.'}), 400
        file = request.files['file']
        level = request.form.get('level')
        print(f"File received: {file.filename}, Level: {level}")

        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)
        print(f"File saved to: {filepath}")

        if file.filename.endswith('.json'):
            with open(filepath, 'r') as f:
                diagram_data = json.load(f)
            print("File parsed as JSON")
        elif file.filename.endswith('.docx'):
            diagram_data = parse_docx_to_c4_json(filepath)
            print("File parsed as DOCX")
        else:
            print("Error: Unsupported file type")
            return jsonify({'error': 'Unsupported file type'}), 400

        if not diagram_data:
            print("Error: Failed to parse document")
            return jsonify({'error': 'Failed to parse document.'}), 500

        diagram = None
        output_filename = f"{level}_diagram"
        print(f"Output filename: {output_filename}")

        if level == "c1":
            diagram = C4ContextDiagram(system_name=diagram_data.get("system_name", "Unnamed System"), output_filename=output_filename)
        elif level == "c2":
            diagram = C4ContainerDiagram(system_name=diagram_data.get("system_name", "Unnamed System"), output_filename=output_filename)
        elif level == "c3":
            diagram = C4ComponentDiagram(container_name=diagram_data.get("container", "Unnamed Container"), output_filename=output_filename)
        elif level == "c4":
            diagram = C4CodeDiagram(component_name=diagram_data.get("component_name", "Unnamed Component"), output_filename=output_filename)

        if diagram:
            diagram.from_json(diagram_data)
            print("Diagram object created and populated")
            return jsonify({
                'message': 'Diagram parsed successfully!',
                'c4Data': diagram.to_json(),
                'level': level
            }), 200
        else:
            print("Error: Failed to create diagram object")
            return jsonify({'error': 'Failed to create diagram object'}), 500

    except Exception as e:
        print(f"Server error: {e}")  # Print the full exception
        return jsonify({'error': f'Server error: {str(e)}'}), 500
@app.route('/diagrams_output/<path:filename>')
def serve_diagram(filename):
    return send_from_directory(DIAGRAM_FOLDER, filename)

if __name__ == "__main__":
    app.run(debug=True)